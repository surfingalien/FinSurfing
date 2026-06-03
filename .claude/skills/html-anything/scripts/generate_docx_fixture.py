#!/usr/bin/env python3
"""Generate examples/docx/input.docx — a synthetic, anonymized memo / RFC
fixture for html-anything.

The shape mimics a real internal product memo: status block, decisions
log, claims with rationale, action items with owners and dates, an open-
questions section. Everything is invented.

Run:  python3 scripts/generate_docx_fixture.py
Output: examples/docx/input.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUT = Path(__file__).resolve().parent.parent / "examples" / "docx" / "input.docx"


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_p(doc, text, italic=False, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x41, 0x38)
    return p


def add_kv_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for cell in t.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True


def add_action_table(doc, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=4)
    t.style = "Light Grid Accent 1"
    headers = ["Action", "Owner", "Due", "Status"]
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10.5)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10.5)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title block ---
    title = doc.add_heading("RFC-014: Pricing Page V2 — Decision Memo", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_p(
        doc,
        "Synthetic fixture for html-anything. All names, numbers, customer "
        "quotes, deal sizes, and internal projects in this memo are invented.",
        italic=True, color=RGBColor(0x8d, 0x71, 0x66), size=10,
    )

    add_kv_table(doc, [
        ("Author", "Riley Chen (synthetic, Product)"),
        ("Reviewers", "Asha Vora (Eng), Marcus Liang (GTM), Devon Park (Finance)"),
        ("Status", "Decided — implementation kickoff Q2 2026"),
        ("Decision date", "2026-04-12"),
        ("Last updated", "2026-04-22"),
        ("Related", "RFC-009 (free-tier kill), PROD-882 (pricing page redesign)"),
    ])

    # --- TL;DR ---
    add_h(doc, "TL;DR", level=1)
    add_p(doc, (
        "We are replacing the current 4-tier pricing page with a 2-tier model "
        "(Starter, Team) plus a contact-sales Enterprise track. The change is "
        "driven by three findings: (1) the Pro tier captured only 7% of new "
        "signups in Q1 but consumed 31% of pricing-page time-on-page, (2) "
        "support tickets tagged 'pricing confusion' rose 2.4x year-over-year, "
        "and (3) every winning competitor in our segment has consolidated to "
        "≤ 3 published tiers. Expected impact: +18% activation rate on the "
        "pricing page, neutral net revenue in 2 quarters, +6% by Q4 2026."
    ))

    # --- Context ---
    add_h(doc, "1. Context", level=1)
    add_p(doc, (
        "The current pricing page (v1, shipped 2024) presents Free, Starter, "
        "Pro, and Team tiers. Free was killed in RFC-009 but the Pro tier "
        "remains — it was originally pitched as the 'power individual' "
        "segment between Starter and Team."
    ))
    add_p(doc, (
        "Three things changed since 2024 that invalidate the original "
        "pitch:"
    ))
    add_p(doc, (
        "First, the prosumer segment we expected to fill Pro never showed "
        "up. Q1 2026 signups split 71% Starter, 7% Pro, 22% Team. Pro is "
        "a quiet middle that almost no one chooses on first contact."
    ))
    add_p(doc, (
        "Second, the support cost of the four-tier table is real. We "
        "logged 412 'pricing confusion' tickets in Q1 2026 versus 173 "
        "in Q1 2025. The most common confusion: customers picked Pro "
        "thinking it included multi-seat features (it does not) and "
        "downgraded within 30 days."
    ))
    add_p(doc, (
        "Third, every competitor with positive growth in our segment "
        "ships ≤ 3 published tiers. The two that grew fastest (Linear, "
        "Vercel) ship 2 + 'enterprise'. The cognitive load argument is "
        "now backed by a market signal."
    ))

    add_quote(doc, (
        "“I picked Pro because it sounded important. Then I realised "
        "it gave me nothing my coworkers could actually use, and I "
        "downgraded the next day.” — synthetic customer interview, "
        "March 2026"
    ))

    # --- Proposal ---
    add_h(doc, "2. Proposal", level=1)
    add_p(doc, (
        "Replace the four-tier table with two published tiers plus a "
        "contact-sales Enterprise option:"
    ))

    add_h(doc, "2.1 Tiers", level=2)
    add_p(doc, "Starter — single user, $19/mo. Existing Starter customers grandfathered.", size=11)
    add_p(doc, "Team — up to 25 seats, $39/seat/mo. Existing Team customers see no price change.", size=11)
    add_p(doc, "Enterprise — contact sales. SSO, audit log, custom DPA, named support.", size=11)

    add_h(doc, "2.2 What happens to Pro customers", level=2)
    add_p(doc, (
        "There are 1,184 active Pro customers. They will be migrated as "
        "follows: solo Pro accounts → grandfathered Starter at current "
        "Pro price for 12 months, then offered Starter or Team upgrade "
        "with a 20% loyalty discount. Multi-seat Pro accounts (these "
        "exist accidentally; the tier was never supposed to allow more "
        "than one seat) → free upgrade to Team at Pro pricing for 6 "
        "months."
    ))

    add_h(doc, "2.3 Page design changes", level=2)
    add_p(doc, (
        "Three-column layout with the Team column visually emphasised "
        "(brand orange border, ‘Most popular’ tag). Feature comparison "
        "table simplified from 47 rows to 19. FAQ section reduced from "
        "23 questions to 8, with the four highest-traffic ones promoted "
        "above the fold."
    ))

    # --- Alternatives ---
    add_h(doc, "3. Alternatives Considered", level=1)
    add_p(doc, (
        "<b>3.1 Keep four tiers but rename Pro.</b> Cheaper to ship, but "
        "leaves the cognitive-load problem unresolved and the migration "
        "risk in place. Rejected because the data shows the problem is "
        "structural, not naming."
    ))
    add_p(doc, (
        "<b>3.2 Single tier + usage-based.</b> Considered seriously. "
        "Strong long-term direction; wrong short-term move because our "
        "billing infrastructure does not yet support metered components "
        "and we are not staffed to build it before Q4. Revisit in 2027."
    ))
    add_p(doc, (
        "<b>3.3 Three published tiers (Starter, Team, Business).</b> "
        "Closest to v1 minus Pro. Rejected because adding a Business "
        "tier reintroduces the 'middle confusion' pattern we are trying "
        "to remove. Two-plus-Enterprise is structurally cleaner."
    ))

    # --- Risks ---
    add_h(doc, "4. Risks & Mitigations", level=1)
    add_p(doc, (
        "<b>4.1 Pro customer churn.</b> The migration plan extends current "
        "pricing for 6–12 months, which we project converts 78% to Starter "
        "or Team and 22% to churn. Worst-case (50% churn) loses ~$680K ARR. "
        "Mitigation: dedicated CSM outreach to the top 100 Pro accounts."
    ))
    add_p(doc, (
        "<b>4.2 Sales pipeline impact.</b> Removing the published Pro tier "
        "may reduce inbound qualification accuracy. Mitigation: run a "
        "qualification quiz on the pricing page (already prototyped in "
        "PROD-851) and route Pro-shaped leads directly to AEs."
    ))
    add_p(doc, (
        "<b>4.3 Reputational risk.</b> Pricing changes draw attention. "
        "Mitigation: ship a public-blog post explaining the change "
        "before the page goes live, with a dedicated migration FAQ."
    ))

    # --- Decisions ---
    add_h(doc, "5. Decisions", level=1)
    add_kv_table(doc, [
        ("D-1", "Adopt 2-tier-plus-Enterprise model. Approved 2026-04-12."),
        ("D-2", "Pro tier sunset in waves: solo accounts first, multi-seat second."),
        ("D-3", "Migration window: 12 months for solo, 6 months for multi-seat."),
        ("D-4", "Loyalty discount: 20% off list for migrating Pro customers, 12 months."),
        ("D-5", "Q1 2027 review for usage-based tier feasibility (links to RFC-014A)."),
    ])

    # --- Action items ---
    add_h(doc, "6. Action Items", level=1)
    add_action_table(doc, [
        ["Update pricing page comp table to 19 rows", "Riley Chen", "2026-05-15", "In progress"],
        ["Migration email sequence + drip", "Marcus Liang", "2026-05-22", "Drafted"],
        ["CSM outreach plan for top 100 Pro accounts", "Devon Park", "2026-05-08", "Done"],
        ["Pricing-page qualification quiz wiring", "Asha Vora", "2026-06-02", "Not started"],
        ["Public blog post + migration FAQ", "Riley Chen", "Pre-launch -7d", "Outline"],
        ["Billing infra spike for usage-based (RFC-014A)", "Asha Vora", "2026-09-30", "Scoped"],
    ])

    # --- Open questions ---
    add_h(doc, "7. Open Questions", level=1)
    add_p(doc, "Q1. Should the loyalty discount apply for 12 or 18 months? Current proposal is 12; Marcus argues 18 to soften the change.")
    add_p(doc, "Q2. Do we publish per-seat Team pricing or 'starting at $39/seat'? Finance prefers the latter; design prefers the former.")
    add_p(doc, "Q3. Enterprise floor — is $24K/year the right minimum? Current AE feedback is the floor is too low and is attracting Team-shaped accounts that should self-serve.")
    add_p(doc, "Q4. Migration FAQ tone — should the post acknowledge that the Pro tier was a strategic miss, or frame the change as a forward-looking simplification?")

    # --- Appendix ---
    add_h(doc, "Appendix A — Data references", level=1)
    add_p(doc, "All numbers in this synthetic memo are invented and intended to mimic the shape of real internal pricing analyses, not the substance.", italic=True, size=10)
    add_p(doc, "A.1 Q1 2026 signup mix: 71% Starter / 7% Pro / 22% Team (n = 14,732 signups, synthetic)")
    add_p(doc, "A.2 Pricing-confusion ticket volume: 412 (Q1 2026) vs 173 (Q1 2025), +138% YoY")
    add_p(doc, "A.3 Pro tier active customers: 1,184 — of which 1,061 solo, 123 multi-seat (synthetic)")
    add_p(doc, "A.4 Projected ARR impact: -$210K (months 1–6), neutral by month 9, +6% baseline by Q4 2026")
    add_p(doc, "A.5 Competitor pricing-tier counts (synthetic survey): Linear 2, Vercel 2+E, Notion 3, Asana 4, Monday 5")

    doc.save(str(OUT))
    print(f"wrote {OUT.relative_to(OUT.parent.parent.parent)} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
